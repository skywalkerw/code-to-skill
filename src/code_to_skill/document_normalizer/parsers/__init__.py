"""文档解析器集合。

Markdown / HTML / PDF / DOCX / Text 解析器。
"""
from __future__ import annotations

import re
from typing import Any

from code_to_skill.document_normalizer.types import RawDocument


# ── Markdown ────────────────────────────────────────────────

def parse_markdown(text: str) -> list[dict]:
    """解析 Markdown 文本为 blocks。"""
    blocks: list[dict] = []
    lines = text.split("\n")
    in_code_block = False
    code_lines: list[str] = []

    for line in lines:
        stripped = line.strip()

        # code fence
        if stripped.startswith("```"):
            if in_code_block:
                blocks.append({"type": "code_block", "text": "\n".join(code_lines), "language": stripped[3:] or ""})
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # heading
        m = re.match(r"^(#{1,6})\s+(.+)", stripped)
        if m:
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2)})
            continue

        # table
        if "|" in stripped and stripped.startswith("|"):
            blocks.append({"type": "table_row", "text": stripped})
            continue

        # list
        if re.match(r"^[\-\*\+]\s+", stripped) or re.match(r"^\d+\.\s+", stripped):
            blocks.append({"type": "list_item", "text": stripped})
            continue

        # paragraph
        if stripped:
            blocks.append({"type": "paragraph", "text": stripped})

    return blocks


# ── HTML ────────────────────────────────────────────────────

def parse_html(text: str) -> list[dict]:
    """解析 HTML 为 blocks。使用 beautifulsoup4（可选）或正则降级。"""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(text, "html.parser")
        blocks: list[dict] = []

        for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "pre", "table", "li"]):
            if tag.name.startswith("h"):
                level = int(tag.name[1])
                blocks.append({"type": "heading", "level": level, "text": tag.get_text(strip=True)})
            elif tag.name == "p":
                text = tag.get_text(strip=True)
                if text:
                    blocks.append({"type": "paragraph", "text": text})
            elif tag.name == "pre":
                blocks.append({"type": "code_block", "text": tag.get_text()})
            elif tag.name == "table":
                blocks.append({"type": "table", "text": str(tag), "caption": ""})
            elif tag.name == "li":
                blocks.append({"type": "list_item", "text": tag.get_text(strip=True)})

        return blocks
    except ImportError:
        return _parse_html_regex(text)


def _parse_html_regex(text: str) -> list[dict]:
    """正则降级 HTML 解析。"""
    blocks: list[dict] = []
    # 移除标签，保留文本
    clean = re.sub(r"<[^>]+>", " ", text)
    clean = re.sub(r"\s+", " ", clean).strip()
    if clean:
        blocks.append({"type": "paragraph", "text": clean})
    return blocks


# ── PDF ─────────────────────────────────────────────────────

def parse_pdf(content: bytes) -> list[dict]:
    """解析 PDF 为 blocks。使用 pdfplumber。"""
    try:
        import pdfplumber
        import io
        blocks: list[dict] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text:
                    for line in text.split("\n"):
                        line = line.strip()
                        if line:
                            blocks.append({
                                "type": "paragraph",
                                "text": line,
                                "page": page_num,
                            })
                # 表格
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        blocks.append({
                            "type": "table",
                            "rows": table,
                            "page": page_num,
                        })
        return blocks
    except ImportError:
        return [{"type": "paragraph", "text": "(pdfplumber not installed, PDF content skipped)"}]


# ── DOCX ────────────────────────────────────────────────────

def parse_docx(content: bytes) -> list[dict]:
    """解析 DOCX 为 blocks。使用 python-docx。"""
    try:
        import io
        from docx import Document
        blocks: list[dict] = []
        doc = Document(io.BytesIO(content))
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                blocks.append({"type": "heading", "level": level, "text": text})
            else:
                blocks.append({"type": "paragraph", "text": text})
        # 表格
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            blocks.append({"type": "table", "rows": rows})
        return blocks
    except ImportError:
        return [{"type": "paragraph", "text": "(python-docx not installed)"}]


# ── Text ────────────────────────────────────────────────────

def parse_text(text: str) -> list[dict]:
    """纯文本降级解析。"""
    return [{"type": "paragraph", "text": line.strip()}
            for line in text.split("\n") if line.strip()]


# ── 统一接口 ────────────────────────────────────────────────

def parse_raw_document(raw: RawDocument) -> list[dict]:
    """根据 RawDocument 类型选择解析器，返回 blocks 列表。"""
    st = raw.source_type
    if st == "markdown":
        return parse_markdown(raw.text)
    elif st == "html":
        return parse_html(raw.text)
    elif st == "pdf":
        return parse_pdf(raw.content)
    elif st == "docx":
        return parse_docx(raw.content)
    else:
        return parse_text(raw.text)
